import praw
from docx import Document

# Initialize the Reddit instance
reddit = praw.Reddit(client_id='WC-c4C5SG4j0x9nVbG_SmQ',
                     client_secret='dRWVFQ6Q7OjR9AMnm9aRuKJQRysbrw',
                     user_agent='AgentHistorical9521')

# Subreddits to target
subreddits = ['cars', 'electricvehicles', 'TeslaMotors', 'BMW', 'Audi', 'Ford', 'environment', 'sustainability', 'ClimateActionPlan', 'technology', 'Futurology', 'science', 'ProductReviews', 'ConsumerReports']

# Keywords to search for
keywords = ['sustainability', 'carbon neutrality', 'electric vehicle', 'EV', 'hybrid', 'renewable energy', 'climate action']

# Create a Word document
doc = Document()
doc.add_heading('Reddit Discussions on Automotive Sustainability', 0)

for subreddit in subreddits:
    subreddit_instance = reddit.subreddit(subreddit)
    doc.add_heading(f'Subreddit: {subreddit}', level=1)
    
    for keyword in keywords:
        doc.add_heading(f'Keyword: {keyword}', level=2)
        
        for submission in subreddit_instance.search(keyword, limit=10):  # Limiting to 10 for brevity
            doc.add_heading(submission.title, level=3)
            doc.add_paragraph(f'Score: {submission.score}')
            doc.add_paragraph(f'Comments: {submission.num_comments}')
            doc.add_paragraph(f'URL: {submission.url}')
            doc.add_paragraph('---')

# Save the document
doc.save('Reddit_Automotive_Sustainability.docx')
print("Data has been saved to 'Reddit_Automotive_Sustainability.docx'")